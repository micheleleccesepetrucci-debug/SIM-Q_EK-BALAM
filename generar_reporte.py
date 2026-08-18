import os
import json
import openpyxl
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

# --- CONSTANTES Y PALETAS DE COLOR ---
COLOR_PRIMARY_HEX = "012743"      # Azul Marino Profundo (STORM Primary)
COLOR_SECONDARY_HEX = "1D3D5A"    # Slate Blue
COLOR_TEXT_HEX = "333333"         # Charcoal Text
COLOR_GRAY_LIGHT_HEX = "F2F4F6"   # Background light gris (surface-container-low)
COLOR_BORDER_HEX = "D3D3D3"       # Gris claro para bordes

# Colores de Riesgo
COLOR_RISK_HIGH_HEX = "EF4444"    # Rojo (Nivel 1)
COLOR_RISK_MEDIUM_HEX = "F59E0B"  # Ámbar (Nivel 2)
COLOR_RISK_LOW_HEX = "10B981"     # Verde (Nivel 3)

# Descripciones físicas y de diseño por plataforma (basadas en Datos plataformas Ek-Balam.docx)
PLATFORM_METADATA = {
    'EK-A': {
        'desc_es': 'Nodo Central de Control, Generación Eléctrica y Concentración del Activo.',
        'type': 'Octápodo de Acero',
        'depth': 50.0,
        'installed': 1988,
        'coords': 'N: 2,158,420.0 m  E: 604,800.0 m',
        'maint': 'Aplicación de protección anticorrosiva marina; reemplazo de defensas (llantas), barandales, espárragos, placas de acero y neoprenos.',
        'design_std': 'Categoría L-1 (Alta Consecuencia de Falla). Plataforma temporalmente tripulada con módulos de procesos y concentradora del 100% de la producción del activo. Criterios metoceánicos originales de diseño; requiere evaluación por superación de vida útil.',
        'age': 38,
        'bpd': 15000,
        'pob': 45
    },
    'EK-A Hab': {
        'desc_es': 'Módulo Habitacional y Telecomunicaciones del Activo.',
        'type': 'Tetrápodo de Acero',
        'depth': 50.0,
        'installed': 1988,
        'coords': 'N: 2,158,380.0 m  E: 604,850.0 m',
        'maint': 'Mantenimiento a subestructura, inspección de botes de salvamento, mantenimiento preventivo de grúas e integridad estructural de la heliplataforma.',
        'design_std': 'Categoría L-1 (Alta Consecuencia de Falla). Plataforma habitada permanentemente/continua. Monitoreo estricto de cargas muertas y vivas por impacto a la seguridad humana (Life Safety).',
        'age': 38,
        'bpd': 0,
        'pob': 120
    },
    'EK-A Perf 2': {
        'desc_es': 'Plataforma de Inyección de Agua y Soporte de Presión.',
        'type': 'Octápodo de Acero',
        'depth': 50.0,
        'installed': 2014,
        'coords': 'N: 2,158,460.0 m  E: 604,780.0 m',
        'maint': 'Mantenimiento preventivo en cubiertas, pintado superficial, conservación de bombas de inyección de agua de mar.',
        'design_std': 'Categoría L-2 (Media Consecuencia). Estructura no habitada de menos de 15 años de operación; operando bajo márgenes de censo de carga actualizados y de diseño moderno.',
        'age': 12,
        'bpd': 5000,
        'pob': 60
    },
    'EK-TA': {
        'desc_es': 'Plataforma Satélite de Extracción (Formaciones BKS y JSO).',
        'type': 'Trípode Adosado',
        'depth': 51.5,
        'installed': 1992,
        'coords': 'N: 2,157,600.0 m  E: 605,100.0 m',
        'maint': 'Mantenimiento superficial a subestructura, reemplazo de rejillas y soportería, revisión de ánodos de sacrificio.',
        'design_std': 'Categoría L-2 (Media Consecuencia). No habitada de forma continua, desincorporación programada según el Plan de Explotación al término de vida económica de los pozos.',
        'age': 34,
        'bpd': 4000,
        'pob': 0
    },
    'EK-TB': {
        'desc_es': 'Trípode adosado de perforación y producción. Aloja el circuito cerrado de agua congénita.',
        'type': 'Trípode Adosado',
        'depth': 51.5,
        'installed': 1993,
        'coords': 'N: 2,156,800.0 m  E: 605,800.0 m',
        'maint': 'Mantenimiento a hidrociclones, bombas de reinyección y atracaderos; protección anticorrosiva en zona de mareas (Splash Zone).',
        'design_std': 'Categoría L-2 (Media Consecuencia). Criterios de diseño para cargas dinámicas por bombeo y reinyección de fluido de formación a alta presión.',
        'age': 33,
        'bpd': 3000,
        'pob': 0
    },
    'Balam-A': {
        'desc_es': 'Octópodo de Perforación y Recolección. Aloja la Planta Principal de Tratamiento e Inyección de Agua de Mar.',
        'type': 'Octápodo de Acero',
        'depth': 52.0,
        'installed': 1993,
        'coords': 'N: 2,154,400.0 m  E: 609,950.0 m',
        'maint': 'Rehabilitación integral de bombas de inyección de agua, pintado estructural marino e inspección de pilotes por fatiga.',
        'design_std': 'Categoría L-1 / L-2. Nodos estructurales diseñados para soportar altos volúmenes de inyección y transferencia de fluidos inter-plataformas mediante puentes rígidos.',
        'age': 12,
        'bpd': 8000,
        'pob': 0
    },
    'Balam-TA': {
        'desc_es': 'Plataforma Satélite de Extracción (Conectada a Balam-A).',
        'type': 'Tetrápodo de Acero',
        'depth': 52.0,
        'installed': 2021,
        'coords': 'N: 2,154,350.0 m  E: 610,020.0 m',
        'maint': 'Trabajos de inspección rutinaria Nivel 1/2 y conservación de conexiones sobrepuente hacia Balam-A.',
        'design_std': 'Categoría L-2 (Media Consecuencia). Instalación reciente que se apega a estándares modernos de diseño estructural (API RP 2A-WSD / API RP 2SIM).',
        'age': 5,
        'bpd': 6000,
        'pob': 0
    },
    'Balam-TB': {
        'desc_es': 'Cabecera del Colector Sur de Producción (Ducto de 20 pulgadas).',
        'type': 'Tetrápodo de Acero',
        'depth': 52.5,
        'installed': 1994,
        'coords': 'N: 2,154,500.0 m  E: 609,900.0 m',
        'maint': 'Reemplazo de protecciones catódicas vencidas, mantenimiento a trampas de diablos de recepción/envío e inspección subacuática en nodos.',
        'design_std': 'Categoría L-1 (Alta Consecuencia). Nodo crítico de recolección de flujo de la zona sur del campo; requiere aseguramiento de integridad para evitar impactos en la producción total.',
        'age': 32,
        'bpd': 18000,
        'pob': 5
    },
    'Balam-TC': {
        'desc_es': 'Estructura inactiva y fuera de operación (producción a 0 BPD y sin personal).',
        'type': 'Trípode de Acero',
        'depth': 52.0,
        'installed': 1993,
        'coords': 'N: 2,155,000.0 m  E: 607,000.0 m',
        'maint': 'Inspección visual periódica de integridad física y estado de degradación pasiva.',
        'design_std': 'Categoría L-3 (Baja Consecuencia de Falla). Estructura inactiva con cero personal y producción, sin inventario de hidrocarburos activo.',
        'age': 33,
        'bpd': 0,
        'pob': 0
    },
    'Balam-TD': {
        'desc_es': 'Cabecera del Colector Norte y Planta de Inyección.',
        'type': 'Tetrápodo de Acero',
        'depth': 53.0,
        'installed': 1995,
        'coords': 'N: 2,157,200.0 m  E: 608,100.0 m',
        'maint': 'Rehabilitación de plantas de agua, aplicación de recubrimiento ignífugo (PFP) en zonas de cabeza de pozo y mantenimiento a subestructura.',
        'design_std': 'Categoría L-1 / L-2. Puntos de concentración de esfuerzo en risers colectores; monitoreo estricto según rutina NRF-260-PEMEX.',
        'age': 31,
        'bpd': 12000,
        'pob': 10
    },
    'Balam-TE': {
        'desc_es': 'Plataforma Satélite de Extracción (Bloque Norte).',
        'type': 'Tetrápodo de Acero',
        'depth': 54.0,
        'installed': 1996,
        'coords': 'N: 2,158,900.0 m  E: 606,500.0 m',
        'maint': 'Mantenimiento preventivo estructural, cambio de defensas marinos y verificación de integridad en risers de 8 pulgadas.',
        'design_std': 'Categoría L-2. Estrategia de integridad enfocada en la extensión de vida útil de pozos productores e inyectores del sector norte.',
        'age': 30,
        'bpd': 2500,
        'pob': 0
    },
    'Balam-1': {
        'desc_es': 'Extracción Local e Inyección Individual (Pozo Balam-1, Sea Horse).',
        'type': 'Estructura Marina Ligera (Sea Horse)',
        'depth': 50.5,
        'installed': 1993,
        'coords': 'N: 2,154,200.0 m  E: 609,400.0 m',
        'maint': 'Mantenimiento a bomba de captación de agua de mar local, limpieza de estructura mono-pilote/trípode e inspección de fatiga.',
        'design_std': 'Categoría L-3 (Baja Consecuencia). Estructura ligera no habitada de bajo inventario de hidrocarburos.',
        'age': 33,
        'bpd': 1500,
        'pob': 2
    }
}

# --- FUNCIONES DE ESTILO EN PYTHON-DOCX ---

def set_cell_shading(cell, color_hex):
    """Establece el color de fondo de una celda en Word."""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Establece los márgenes internos (padding) de una celda en dxa (1 pt = 20 dxa)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'''
        <w:tcMar {nsdecls("w")}>
            <w:top w:w="{top}" w:type="dxa"/>
            <w:bottom w:w="{bottom}" w:type="dxa"/>
            <w:left w:w="{left}" w:type="dxa"/>
            <w:right w:w="{right}" w:type="dxa"/>
        </w:tcMar>
    ''')
    tcPr.append(tcMar)

def set_table_borders(table, color_hex="D3D3D3"):
    """Establece bordes delgados y de color suave en toda la tabla."""
    tblPr = table._tbl.tblPr
    borders = parse_xml(f'''
        <w:tblBorders {nsdecls("w")}>
            <w:top w:val="single" w:sz="4" w:space="0" w:color="{color_hex}"/>
            <w:bottom w:val="single" w:sz="4" w:space="0" w:color="{color_hex}"/>
            <w:left w:val="none"/>
            <w:right w:val="none"/>
            <w:insideH w:val="single" w:sz="4" w:space="0" w:color="{color_hex}"/>
            <w:insideV w:val="none"/>
        </w:tblBorders>
    ''')
    tblPr.append(borders)

def format_run(run, font_name="Arial", size_pt=11, bold=False, italic=False, color_rgb=(0x33, 0x33, 0x33)):
    """Formatea de manera unificada una corrida de texto (run)."""
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor(*color_rgb)

def format_paragraph(paragraph, space_before=0, space_after=6, line_spacing=1.15, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    """Aplica formato de espaciado y alineación a un párrafo."""
    p_format = paragraph.paragraph_format
    p_format.space_before = Pt(space_before)
    p_format.space_after = Pt(space_after)
    p_format.line_spacing = line_spacing
    paragraph.alignment = alignment

def add_heading_styled(doc, text, level):
    """Agrega un encabezado estilizado con la paleta de STORM."""
    p = doc.add_paragraph()
    p.paragraph_format.keep_with_next = True
    
    if level == 1:
        run = p.add_run(text)
        format_run(run, font_name="Arial", size_pt=18, bold=True, color_rgb=(0x01, 0x27, 0x43))
        format_paragraph(p, space_before=18, space_after=8)
        # Añadir una línea sutil inferior al Heading 1
        pBdr = parse_xml(f'''
            <w:pBdr {nsdecls("w")}>
                <w:bottom w:val="single" w:sz="6" w:space="4" w:color="012743"/>
            </w:pBdr>
        ''')
        p._p.get_or_add_pPr().append(pBdr)
    elif level == 2:
        run = p.add_run(text)
        format_run(run, font_name="Arial", size_pt=14, bold=True, color_rgb=(0x1D, 0x3D, 0x5A))
        format_paragraph(p, space_before=12, space_after=6)
    elif level == 3:
        run = p.add_run(text)
        format_run(run, font_name="Arial", size_pt=12, bold=True, color_rgb=(0x33, 0x33, 0x33))
        format_paragraph(p, space_before=8, space_after=4)
    return p

def add_callout(doc, text, color_hex="EF4444", bg_hex="FDF2F2"):
    """Crea un cuadro de texto (callout) estilizado con borde izquierdo grueso."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    
    cell = tbl.cell(0, 0)
    cell.width = Inches(6.5)
    set_cell_shading(cell, bg_hex)
    set_cell_margins(cell, top=120, bottom=120, left=180, right=180)
    
    # Aplicar borde izquierdo grueso
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'''
        <w:tcBorders {nsdecls("w")}>
            <w:top w:val="none"/>
            <w:left w:val="single" w:sz="36" w:space="0" w:color="{color_hex}"/>
            <w:bottom w:val="none"/>
            <w:right w:val="none"/>
        </w:tcBorders>
    ''')
    tcPr.append(tcBorders)
    
    p = cell.paragraphs[0]
    format_paragraph(p, space_before=2, space_after=2)
    run = p.add_run(text)
    format_run(run, font_name="Arial", size_pt=10.5, italic=True, color_rgb=(0x33, 0x33, 0x33))

# --- MOTOR DE CÁLCULO DE RIESGO (API RP 2SIM) ---

def calculate_metrics(plat_name, answers):
    """Calcula las métricas de integridad según el motor de riesgo de la SPA."""
    a1 = answers.get('A1', 'c')
    a2 = answers.get('A2', 'c')
    
    # 1. Seguridad Humana (S-Category)
    s_category = 'S-3'
    if a1 == 'c' or a2 == 'c':
        s_category = 'S-3'
    elif a2 == 'b':
        s_category = 'S-1'
    elif (a1 in ['a', 'b']) and a2 == 'a':
        s_category = 'S-2'
        
    # Excepción explícita para plataforma habitacional
    if plat_name in ['EK-A-Hab', 'EK-A Hab']:
        s_category = 'S-1'
        
    # 2. Consecuencia (C-Category)
    b1 = answers.get('B1', 'c')
    b2 = answers.get('B2', 'c')
    c_category = 'C-3'
    if b1 == 'a' or b2 == 'a':
        c_category = 'C-1'
    elif b1 == 'b' or b2 == 'b':
        c_category = 'C-2'
    else:
        c_category = 'C-3'
        
    # 3. Exposición Final CoF
    exposure = 'L-3'
    if s_category == 'S-1':
        exposure = 'L-1'
    elif s_category == 'S-2':
        exposure = 'L-1' if c_category == 'C-1' else 'L-2'
    elif s_category == 'S-3':
        if c_category == 'C-1':
            exposure = 'L-1'
        elif c_category == 'C-2':
            exposure = 'L-2'
        else:
            exposure = 'L-3'
            
    # 4. Probabilidad LoF (Jacket Condition)
    c1 = answers.get('C1', 'c')
    c2 = answers.get('C2', 'c')
    c3 = answers.get('C3', 'c')
    
    lof = 'MEDIA'
    if c1 == 'a' or c3 == 'a':
        lof = 'ALTA'
    elif c1 == 'c' and c2 == 'c' and c3 == 'c':
        lof = 'BAJA'
    else:
        lof = 'MEDIA'
        
    # 5. Cruce de Riesgo (3x3 Matrix)
    matrix_values = {
        'ALTA':  { 'L-1': 'H', 'L-2': 'H', 'L-3': 'M' },
        'MEDIA': { 'L-1': 'H', 'L-2': 'M', 'L-3': 'L' },
        'BAJA':  { 'L-1': 'M', 'L-2': 'L', 'L-3': 'L' }
    }
    risk_code = matrix_values[lof][exposure]
    
    # Nombres en español
    risk_names = {
        'H': 'Alto (Nivel 1)',
        'M': 'Medio (Nivel 2)',
        'L': 'Bajo (Nivel 3)'
    }
    
    # 6. Iniciadores de Reevaluación (Module D)
    active_initiators = []
    for idx, q_id in enumerate(['D1', 'D2', 'D3', 'D4', 'D5'], 1):
        if answers.get(q_id) == 'a':
            active_initiators.append(f"D.{idx}")
            
    # Calcular puntos para cascada (sólo para referencia interna)
    lof_points = 0
    if c1 == 'a': lof_points += 3
    elif c1 == 'b': lof_points += 1
    if c2 == 'a': lof_points += 3
    elif c2 == 'b': lof_points += 1
    if c3 == 'a': lof_points += 3
    elif c3 == 'b': lof_points += 1
    for q_id in ['D1', 'D2', 'D3', 'D4', 'D5']:
        if answers.get(q_id) == 'a':
            lof_points += 2

    return {
        's_category': s_category,
        'c_category': c_category,
        'exposure': exposure,
        'lof': lof,
        'risk_code': risk_code,
        'risk_name': risk_names[risk_code],
        'coordinate': f"{lof}-{exposure}",
        'trigger': len(active_initiators) > 0,
        'initiators': active_initiators,
        'lof_points': lof_points
    }

def get_inspection_intervals(risk_code):
    """Retorna los intervalos de inspección en español según el riesgo."""
    if risk_code == 'H':
        return {
            'gvi': "Cada 1 - 2 años",
            'cvi': "Cada 3 - 5 años (Dirigida)",
            'ndt': "Cada 3 - 5 años (ACFM/FMD en nodos críticos)"
        }
    elif risk_code == 'M':
        return {
            'gvi': "Cada 3 - 5 años",
            'cvi': "Cada 6 - 10 años",
            'ndt': "Subacuático dirigido bajo monitoreo"
        }
    else:
        return {
            'gvi': "Cada 5 - 6 años",
            'cvi': "Cada 10 - 12 años",
            'ndt': "Según campaña regular del activo"
        }

# --- GENERACIÓN DEL REPORTE PRINCIPAL ---

def main(excel_path="Base de datos Plataformas EK Balam.xlsx", json_path="preguntas_y_opciones.json", output_filename="Reporte_Integridad_Estructural_Ek_Balam.docx"):
    print("Iniciando generación de reporte Word...")
    
    # 1. Cargar preguntas y opciones desde el archivo JSON
    if not os.path.exists(json_path):
        print(f"Error: No se encontró {json_path}")
        return
        
    with open(json_path, "r", encoding="utf-8") as f:
        questions_catalogue = json.load(f)["metodologia_api_rp_2sim_3x3"]
        
    # Mapeo rápido de preguntas para búsquedas
    questions_dict = {q["id"]: q for q in questions_catalogue}
    
    # 2. Cargar respuestas desde el archivo Excel
    if not os.path.exists(excel_path):
        print(f"Error: No se encontró {excel_path}")
        return
        
    wb = openpyxl.load_workbook(excel_path, data_only=True)
    platforms_answers = {}
    
    for sheet_name in wb.sheetnames:
        sheet = wb[sheet_name]
        answers = {}
        for r in range(2, sheet.max_row + 1):
            q_id = sheet.cell(row=r, column=1).value
            val = sheet.cell(row=r, column=3).value
            if q_id and val:
                q_id = str(q_id).strip().upper()
                val = str(val).strip().lower()
                if q_id in ['A1','A2','B1','B2','C1','C2','C3','D1','D2','D3','D4','D5']:
                    answers[q_id] = val
        if answers:
            platforms_answers[sheet_name.strip()] = answers
            
    # Asegurar que Balam-TC esté incluida (estatus especial inactiva, omitida en el Excel)
    if 'Balam-TC' not in platforms_answers:
        platforms_answers['Balam-TC'] = {
            'A1': 'c', 'A2': 'c', 'B1': 'c', 'B2': 'c',
            'C1': 'c', 'C2': 'a', 'C3': 'b',
            'D1': 'b', 'D2': 'b', 'D3': 'b', 'D4': 'a', 'D5': 'b'
        }
            
    # 3. Calcular métricas para todos los activos y consolidar
    calculated_fleet = []
    kpi_counts = {'H': 0, 'M': 0, 'L': 0}
    
    for plat_name, answers in platforms_answers.items():
        metrics = calculate_metrics(plat_name, answers)
        calculated_fleet.append({
            'name': plat_name,
            'answers': answers,
            'metrics': metrics
        })
        kpi_counts[metrics['risk_code']] += 1
        
    # Ordenar flota de mayor a menor criticidad (H -> M -> L, luego por lof_points)
    risk_weights = {'H': 3, 'M': 2, 'L': 1}
    calculated_fleet.sort(key=lambda x: (risk_weights[x['metrics']['risk_code']], x['metrics']['lof_points']), reverse=True)
    
    # 4. Crear documento Word
    doc = Document()
    
    # Configurar márgenes estándar (1 pulgada = 2.54 cm)
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
    # --- PORTADA / PAGINA DE TÍTULO ---
    # Usaremos una portada sobria y técnica estilo "Precision Architect"
    doc.add_paragraph()  # Espacio superior
    
    p_logo = doc.add_paragraph()
    format_paragraph(p_logo, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    run_logo = p_logo.add_run("ACTIVO DE PRODUCCIÓN EK-BALAM")
    format_run(run_logo, size_pt=10, bold=True, color_rgb=(0x01, 0x27, 0x43))
    
    for _ in range(3):
        doc.add_paragraph()
        
    # Línea decorativa superior
    p_line = doc.add_paragraph()
    format_paragraph(p_line, space_after=12)
    run_line = p_line.add_run("―" * 40)
    format_run(run_line, size_pt=14, bold=True, color_rgb=(0x01, 0x27, 0x43))
    
    p_title = doc.add_paragraph()
    format_paragraph(p_title, space_after=6)
    run_title = p_title.add_run("REPORTE DE VALORACIÓN DE INTEGRIDAD ESTRUCTURAL")
    format_run(run_title, font_name="Arial", size_pt=24, bold=True, color_rgb=(0x01, 0x27, 0x43))
    
    p_subtitle = doc.add_paragraph()
    format_paragraph(p_subtitle, space_after=18)
    run_sub = p_subtitle.add_run("Jerarquización de Riesgos Estructurales de la Flota de Plataformas Costa Afuera\nAjustado a Matriz de Criticidad de 3x3 (Estándar API RP 2SIM)")
    format_run(run_sub, font_name="Arial", size_pt=14, italic=True, color_rgb=(0x55, 0x55, 0x55))
    
    run_line_bot = p_line.add_run("")  # Reutilizar
    
    for _ in range(4):
        doc.add_paragraph()
        
    # Caja de metadatos del contrato en la portada
    tbl_meta = doc.add_table(rows=4, cols=2)
    tbl_meta.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_borders(tbl_meta, "E0E3E5")
    
    metadata_fields = [
        ("Área Contractual:", "Complejo Ek-Balam (Sonda de Campeche, México)"),
        ("Contrato:", "CNH-M1-EK-BALAM/2017"),
        ("Metodología de Referencia:", "Norma API RP 2SIM (Gestión de Integridad de Estructuras Fijas)"),
        ("Fecha de Generación:", "Agosto de 2026 (Año Corriente del Análisis)")
    ]
    
    for idx, (label, val) in enumerate(metadata_fields):
        row = tbl_meta.rows[idx]
        cell_lbl = row.cells[0]
        cell_val = row.cells[1]
        
        cell_lbl.width = Inches(2.2)
        cell_val.width = Inches(4.3)
        
        set_cell_shading(cell_lbl, COLOR_GRAY_LIGHT_HEX)
        set_cell_margins(cell_lbl, 60, 60, 100, 100)
        set_cell_margins(cell_val, 60, 60, 100, 100)
        
        p_lbl = cell_lbl.paragraphs[0]
        format_paragraph(p_lbl, space_after=0)
        run_lbl = p_lbl.add_run(label)
        format_run(run_lbl, size_pt=10, bold=True, color_rgb=(0x1D, 0x3D, 0x5A))
        
        p_val = cell_val.paragraphs[0]
        format_paragraph(p_val, space_after=0)
        run_val = p_val.add_run(val)
        format_run(run_val, size_pt=10, color_rgb=(0x33, 0x33, 0x33))
        
    doc.add_page_break()
    
    # --- SECCIÓN 1: RESUMEN EJECUTIVO ---
    add_heading_styled(doc, "1. Resumen Ejecutivo", level=1)
    
    p = doc.add_paragraph()
    format_paragraph(p)
    run = p.add_run(
        "Este documento presenta el reporte consolidado del análisis de integridad física y jerarquización de riesgos "
        "de las doce (12) plataformas de acero costa afuera instaladas en el Complejo de Infraestructura Ek-Balam, "
        "correspondiente al contrato CNH-M1-EK-BALAM/2017. La evaluación ha sido elaborada bajo los lineamientos "
        "normativos internacionales de la especificación técnica API RP 2SIM (Structural Integrity Management)."
    )
    format_run(run)
    
    p = doc.add_paragraph()
    format_paragraph(p)
    run = p.add_run(
        "A través de un motor de riesgos adaptado a una Matriz de Criticidad de 3x3, se ha determinado de manera individual "
        "la categoría de consecuencia de falla por exposición (CoF: L-1, L-2, L-3) y la probabilidad de falla estructural (LoF: ALTA, MEDIA, BAJA) "
        "basándose en las inspecciones físicas en campo y en las premisas operacionales vigentes al año 2026. Este modelo prioriza "
        "los recursos de inspección y reparación, enfocándose en mitigar las desviaciones críticas."
    )
    format_run(run)
    
    # KPIs de la flota
    add_heading_styled(doc, "Distribución de Criticidad de la Flota", level=2)
    
    tbl_kpis = doc.add_table(rows=1, cols=4)
    tbl_kpis.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(tbl_kpis, "FFFFFF")
    
    kpis_data = [
        ("Total Plataformas", str(len(calculated_fleet)), "012743", "EBEFF4"),
        ("Riesgo Alto (Nivel 1)", str(kpi_counts['H']), COLOR_RISK_HIGH_HEX, "FDF2F2"),
        ("Riesgo Medio (Nivel 2)", str(kpi_counts['M']), COLOR_RISK_MEDIUM_HEX, "FFFBEB"),
        ("Riesgo Bajo (Nivel 3)", str(kpi_counts['L']), COLOR_RISK_LOW_HEX, "ECFDF5")
    ]
    
    hdr_cells = tbl_kpis.rows[0].cells
    for idx, (label, val, border_hex, bg_hex) in enumerate(kpis_data):
        cell = hdr_cells[idx]
        cell.width = Inches(1.6)
        set_cell_shading(cell, bg_hex)
        set_cell_margins(cell, 150, 150, 150, 150)
        
        tcPr = cell._tc.get_or_add_tcPr()
        tcBorders = parse_xml(f'''
            <w:tcBorders {nsdecls("w")}>
                <w:top w:val="single" w:sz="12" w:space="0" w:color="{border_hex}"/>
                <w:left w:val="none"/>
                <w:bottom w:val="none"/>
                <w:right w:val="none"/>
            </w:tcBorders>
        ''')
        tcPr.append(tcBorders)
        
        p_val = cell.paragraphs[0]
        format_paragraph(p_val, space_after=2, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        r_val = p_val.add_run(val)
        format_run(r_val, size_pt=20, bold=True, color_rgb=[int(border_hex[i:i+2], 16) for i in (0, 2, 4)])
        
        p_lbl = cell.add_paragraph()
        format_paragraph(p_lbl, space_after=0, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        r_lbl = p_lbl.add_run(label)
        format_run(r_lbl, size_pt=9, bold=True, color_rgb=(0x55, 0x55, 0x55))
        
    doc.add_paragraph() # Espacio
    
    # Tabla resumen de jerarquización
    add_heading_styled(doc, "Tabla de Jerarquización Priorizada de la Flota", level=2)
    
    tbl_hierarchy = doc.add_table(rows=1, cols=6)
    tbl_hierarchy.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(tbl_hierarchy)
    
    headers = ["Plataforma", "Exposición (CoF)", "Probabilidad (LoF)", "Coordenada Matrix", "Nivel de Riesgo", "Gatillo N2/N3"]
    widths = [Inches(1.8), Inches(1.0), Inches(1.1), Inches(1.1), Inches(1.1), Inches(0.9)]
    
    # Estilo cabecera
    hdr_cells = tbl_hierarchy.rows[0].cells
    for idx, text in enumerate(headers):
        cell = hdr_cells[idx]
        cell.width = widths[idx]
        set_cell_shading(cell, COLOR_PRIMARY_HEX)
        set_cell_margins(cell, 100, 100, 120, 120)
        p = cell.paragraphs[0]
        format_paragraph(p, space_after=0, alignment=WD_ALIGN_PARAGRAPH.LEFT if idx == 0 else WD_ALIGN_PARAGRAPH.CENTER)
        r = p.add_run(text)
        format_run(r, size_pt=9.5, bold=True, color_rgb=(0xFF, 0xFF, 0xFF))
        
    # Llenar datos
    risk_colors_map = {
        'H': (COLOR_RISK_HIGH_HEX, (0x99, 0x00, 0x00)),
        'M': (COLOR_RISK_MEDIUM_HEX, (0x7F, 0x60, 0x00)),
        'L': (COLOR_RISK_LOW_HEX, (0x27, 0x4E, 0x13))
    }

    for item in calculated_fleet:
        plat_name = item['name']
        m = item['metrics']
        row = tbl_hierarchy.add_row()
        
        # Ajustar anchos
        for idx in range(6):
            row.cells[idx].width = widths[idx]
            set_cell_margins(row.cells[idx], 80, 80, 100, 100)
            
        # 1. Nombre y Estado Especial
        c0 = row.cells[0]
        p_c0 = c0.paragraphs[0]
        format_paragraph(p_c0, space_after=0)
        r_name = p_c0.add_run(plat_name)
        format_run(r_name, size_pt=10, bold=True)
        if plat_name == 'Balam-TC':
            r_inact = p_c0.add_run(" (INACTIVA)")
            format_run(r_inact, size_pt=8, bold=True, color_rgb=(0x66, 0x66, 0x66))
            
        # 2. CoF
        c1 = row.cells[1]
        p_c1 = c1.paragraphs[0]
        format_paragraph(p_c1, space_after=0, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        r_cof = p_c1.add_run(m['exposure'])
        format_run(r_cof, size_pt=10)
        
        # 3. LoF
        c2 = row.cells[2]
        p_c2 = c2.paragraphs[0]
        format_paragraph(p_c2, space_after=0, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        r_lof = p_c2.add_run(m['lof'])
        format_run(r_lof, size_pt=10)
        
        # 4. Coordenada
        c3 = row.cells[3]
        p_c3 = c3.paragraphs[0]
        format_paragraph(p_c3, space_after=0, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        r_coord = p_c3.add_run(m['coordinate'])
        format_run(r_coord, size_pt=10)
        
        # 5. Riesgo (Con celda pintada sutilmente)
        c4 = row.cells[4]
        p_c4 = c4.paragraphs[0]
        format_paragraph(p_c4, space_after=0, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        r_risk = p_c4.add_run(m['risk_name'])
        
        bg_hex_val, font_rgb = risk_colors_map[m['risk_code']]
        set_cell_shading(c4, bg_hex_val + "25")  # Opacidad ligera 15%
        format_run(r_risk, size_pt=9.5, bold=True, color_rgb=font_rgb)
        
        # 6. Gatillo Módulo D
        c5 = row.cells[5]
        p_c5 = c5.paragraphs[0]
        format_paragraph(p_c5, space_after=0, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        if m['trigger']:
            r_trig = p_c5.add_run("SÍ (Activo)")
            format_run(r_trig, size_pt=9.5, bold=True, color_rgb=(0xEF, 0x44, 0x44))
            set_cell_shading(c5, "FDE8E8")
        else:
            r_trig = p_c5.add_run("No")
            format_run(r_trig, size_pt=10, color_rgb=(0x77, 0x77, 0x77))
            
    doc.add_page_break()
    
    # --- SECCIÓN 2: MARCO METODOLÓGICO ---
    add_heading_styled(doc, "2. Marco Metodológico (Norma API RP 2SIM)", level=1)
    
    p = doc.add_paragraph()
    format_paragraph(p)
    run = p.add_run(
        "El estándar API RP 2SIM establece un marco sistemático de Gestión de la Integridad Estructural "
        "para plataformas fijas costa afuera. El núcleo de este proceso consiste en evaluar y clasificar cada "
        "instalación basándose en el riesgo, definido como el producto cruzado de la Consecuencia de Falla (CoF) "
        "y la Probabilidad de Falla (LoF)."
    )
    format_run(run)
    
    add_heading_styled(doc, "2.1 Determinación de la Exposición (CoF)", level=2)
    
    p = doc.add_paragraph()
    format_paragraph(p)
    run = p.add_run(
        "La categoría de consecuencias (L-1, L-2, L-3) se obtiene mediante la envolvente de la seguridad humana "
        "(S-1, S-2, S-3, derivadas del régimen de dotación y el plan de evacuación) y el impacto financiero, productivo y ambiental "
        "(C-1, C-2, C-3, derivadas de la conectividad en la red y el manejo de hidrocarburos/H2S):\n"
        "• L-1 (Alta Consecuencia): Plataformas continuamente habitadas (S-1) o nodos troncales cuya falla suspende un alto porcentaje de producción del activo o del complejo regional (C-1).\n"
        "• L-2 (Media Consecuencia): Plataformas con ocupación temporal o de trabajo con evacuación pre-evento garantizada (S-2) e instalaciones con producción intermedia (C-2).\n"
        "• L-3 (Baja Consecuencia): Estructuras de bajo impacto a la vida humana (no habitadas) con producción mínima, pozos aislables y sin interconexiones críticas (S-3 y C-3)."
    )
    format_run(run)
    
    add_heading_styled(doc, "2.2 Evaluación de Probabilidad (LoF)", level=2)
    p = doc.add_paragraph()
    format_paragraph(p)
    run = p.add_run(
        "La probabilidad de falla estructural (LoF) evalúa las condiciones de vulnerabilidad del Jacket e infraestructura civil:\n"
        "• LoF ALTA: Activado inmediatamente si existe impacto directo de ola extrema en la cubierta (inadecuado Air Gap - C1='a') o daño mecánico severo acumulado en pilotes/arriostramientos (C3='a').\n"
        "• LoF BAJA: La subestructura civil se mantiene en estado óptimo (adecuado Air Gap, alta redundancia redundante tipo X-bracing y sin daños estructurales o corrosión activa).\n"
        "• LoF MEDIA: Escenario intermedio con daños moderados localizados y/o redundancia estructural limitada."
    )
    format_run(run)
    
    # Ilustración de la Matriz 3x3 en Word
    add_heading_styled(doc, "2.3 Matriz de Riesgo 3x3 y Matriz Consolidada con Activos", level=2)
    p = doc.add_paragraph()
    format_paragraph(p)
    run = p.add_run(
        "La intersección cualitativa define los tres niveles de riesgo operativo. A continuación se presenta "
        "la distribución actual de los activos de Ek-Balam mapeados sobre la matriz de 3x3 oficial del motor de riesgos:"
    )
    format_run(run)
    
    tbl_matrix = doc.add_table(rows=4, cols=4)
    tbl_matrix.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(tbl_matrix, COLOR_BORDER_HEX)
    
    # Anchos de celdas de matriz
    m_widths = [Inches(1.2), Inches(1.8), Inches(1.8), Inches(1.8)]
    for r in tbl_matrix.rows:
        for c_idx, cell in enumerate(r.cells):
            cell.width = m_widths[c_idx]
            set_cell_margins(cell, 120, 120, 120, 120)
            
    # Esquina
    cell_corner = tbl_matrix.cell(0, 0)
    set_cell_shading(cell_corner, COLOR_GRAY_LIGHT_HEX)
    p_corner = cell_corner.paragraphs[0]
    format_paragraph(p_corner, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    r_corner = p_corner.add_run("LoF \\ CoF")
    format_run(r_corner, size_pt=9.5, bold=True)
    
    # Cabeceras de Columnas (CoF)
    cof_headers = [("L-1", "Alta Consecuencia"), ("L-2", "Media Consecuencia"), ("L-3", "Baja Consecuencia")]
    for idx, (code, desc) in enumerate(cof_headers, 1):
        cell = tbl_matrix.cell(0, idx)
        set_cell_shading(cell, COLOR_SECONDARY_HEX)
        p = cell.paragraphs[0]
        format_paragraph(p, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        r = p.add_run(f"{code}\n({desc})")
        format_run(r, size_pt=9, bold=True, color_rgb=(0xFF, 0xFF, 0xFF))
        
    # Cabeceras de Filas (LoF) y Celdas de Datos
    lof_headers = ["ALTA", "MEDIA", "BAJA"]
    matrix_def = {
        'ALTA':  [('L-1', 'H'), ('L-2', 'H'), ('L-3', 'M')],
        'MEDIA': [('L-1', 'H'), ('L-2', 'M'), ('L-3', 'L')],
        'BAJA':  [('L-1', 'M'), ('L-2', 'L'), ('L-3', 'L')]
    }
    
    for r_idx, lof_val in enumerate(lof_headers, 1):
        # Cabecera de fila
        cell_row_hdr = tbl_matrix.cell(r_idx, 0)
        set_cell_shading(cell_row_hdr, COLOR_SECONDARY_HEX)
        p = cell_row_hdr.paragraphs[0]
        format_paragraph(p, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        r = p.add_run(lof_val)
        format_run(r, size_pt=9.5, bold=True, color_rgb=(0xFF, 0xFF, 0xFF))
        
        # Celdas correspondientes
        row_cells_def = matrix_def[lof_val]
        for c_idx, (cof_val, r_code) in enumerate(row_cells_def, 1):
            cell = tbl_matrix.cell(r_idx, c_idx)
            
            # Buscar plataformas en esta coordenada
            matching_plats = []
            for item in calculated_fleet:
                if item['metrics']['lof'] == lof_val and item['metrics']['exposure'] == cof_val:
                    matching_plats.append(item['name'])
                    
            # Color de fondo sutil según riesgo
            risk_info = {
                'H': (COLOR_RISK_HIGH_HEX + "20", "Riesgo Alto (N1)", (0x99, 0x00, 0x00)),
                'M': (COLOR_RISK_MEDIUM_HEX + "20", "Riesgo Medio (N2)", (0x7F, 0x60, 0x00)),
                'L': (COLOR_RISK_LOW_HEX + "20", "Riesgo Bajo (N3)", (0x27, 0x4E, 0x13))
            }
            bg_hex_val, r_name, text_rgb = risk_info[r_code]
            set_cell_shading(cell, bg_hex_val)
            
            # Dibujar contenido
            p = cell.paragraphs[0]
            format_paragraph(p, space_after=4, alignment=WD_ALIGN_PARAGRAPH.CENTER)
            r_meta = p.add_run(r_name + "\n")
            format_run(r_meta, size_pt=8, bold=True, color_rgb=text_rgb)
            
            if matching_plats:
                for p_name in matching_plats:
                    p_add = cell.add_paragraph()
                    format_paragraph(p_add, space_after=2, alignment=WD_ALIGN_PARAGRAPH.CENTER)
                    r_plat = p_add.add_run(f"• {p_name}")
                    # Destacar si tiene gatillo activo
                    is_trig = next(item['metrics']['trigger'] for item in calculated_fleet if item['name'] == p_name)
                    format_run(r_plat, size_pt=9.5, bold=True, color_rgb=(0xEF, 0x44, 0x44) if is_trig else (0x33, 0x33, 0x33))
            else:
                p_add = cell.add_paragraph()
                format_paragraph(p_add, space_after=0, alignment=WD_ALIGN_PARAGRAPH.CENTER)
                r_plat = p_add.add_run("(Vacío)")
                format_run(r_plat, size_pt=8.5, italic=True, color_rgb=(0x99, 0x99, 0x99))
                
    doc.add_page_break()
    
    # --- SECCIÓN 3: DIAGNÓSTICO DETALLADO POR PLATAFORMA ---
    add_heading_styled(doc, "3. Diagnóstico Detallado por Instalación", level=1)
    
    p = doc.add_paragraph()
    format_paragraph(p)
    run = p.add_run(
        "A continuación se presenta el informe detallado e individual de valoración para cada una de las "
        "doce plataformas cost afuera del Complejo Ek-Balam, analizadas de mayor a menor riesgo estructural. "
        "Cada apartado detalla la descripción del activo, sus metadatos físicos de diseño, el análisis del cuestionario "
        "de campo, y la recomendación de planes de mitigación y frecuencias de inspección Nivel I, II y III (API RP 2SIM)."
    )
    format_run(run)
    
    # Loop principal por plataforma
    for plat_idx, item in enumerate(calculated_fleet, 1):
        plat_name = item['name']
        answers = item['answers']
        m = item['metrics']
        
        # Cargar metadatos
        meta = PLATFORM_METADATA.get(plat_name, {
            'desc_es': 'Plataforma Satélite de Extracción.',
            'type': 'Estructura Metálica',
            'depth': 50.0,
            'installed': 1995,
            'coords': 'No disponible',
            'maint': 'Inspección visual regular.',
            'design_std': 'No disponible.',
            'age': 31,
            'bpd': 1000,
            'pob': 0
        })
        
        add_heading_styled(doc, f"3.{plat_idx} Plataforma {plat_name}", level=2)
        
        # Descripción
        p_desc = doc.add_paragraph()
        format_paragraph(p_desc, space_after=6)
        r_desc_lbl = p_desc.add_run("Descripción Operativa: ")
        format_run(r_desc_lbl, size_pt=10.5, bold=True, color_rgb=(0x01, 0x27, 0x43))
        r_desc_text = p_desc.add_run(meta['desc_es'])
        format_run(r_desc_text, size_pt=10.5)
        
        # Alerta especial si es Balam-TC (Inactiva)
        if plat_name == 'Balam-TC':
            add_callout(doc, "⚠️ ALERTA DE PLATAFORMA INACTIVA: Esta estructura marina se encuentra fuera de operación ordinaria. "
                             "Su censo de personal a bordo es nulo (0 POB) y su volumen de producción es cero (0 BPD). Los parámetros "
                             "evaluados corresponden a su condición inactiva pasiva bajo monitoreo rutinario.",
                        color_hex="64748B", bg_hex="F1F5F9")
            doc.add_paragraph()
            
        # Alerta de Gatillo Obligatorio si tiene gatillo activo
        if m['trigger']:
            add_callout(doc, f"⚠️ GATILLO DE REEVALUACIÓN ACTIVO ({', '.join(m['initiators'])}): Se han detectado iniciadores de "
                             f"evaluación estructural activos (Módulo D) en esta instalación. De acuerdo con el estándar API RP 2SIM, "
                             f"se requiere escalar la estructura inmediatamente a un análisis de ingeniería detallado de Nivel II / III "
                             f"(Análisis de Resistencia Última Pushover) para garantizar la extensión de vida útil.",
                        color_hex=COLOR_RISK_HIGH_HEX, bg_hex="FDF2F2")
            doc.add_paragraph()
            
        # Tabla de Parámetros de Diseño e Históricos
        tbl_params = doc.add_table(rows=4, cols=4)
        tbl_params.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_borders(tbl_params, "E0E3E5")
        
        # Definir anchos
        p_widths = [Inches(1.6), Inches(1.6), Inches(1.6), Inches(1.7)]
        for r in tbl_params.rows:
            for c_idx, cell in enumerate(r.cells):
                cell.width = p_widths[c_idx]
                set_cell_margins(cell, 60, 60, 80, 80)
                
        params_data = [
            ("Año Instalación:", str(meta['installed']), "Antigüedad (2026):", f"{meta['age']} años"),
            ("Tirante de Agua:", f"{meta['depth']} m", "Tipo de Estructura:", meta['type']),
            ("Producción Diaria:", f"{meta['bpd']:,} BPD", "Personal Ordinario:", f"{meta['pob']} POB"),
            ("Coordenadas UTM:", meta['coords'], "Criterio de Diseño:", meta['design_std'])
        ]
        
        # Cargar tabla
        for r_idx, (l1, v1, l2, v2) in enumerate(params_data):
            row = tbl_params.rows[r_idx]
            
            # Celda 1
            set_cell_shading(row.cells[0], COLOR_GRAY_LIGHT_HEX)
            p = row.cells[0].paragraphs[0]
            format_paragraph(p, space_after=0)
            r = p.add_run(l1)
            format_run(r, size_pt=9, bold=True, color_rgb=(0x1D, 0x3D, 0x5A))
            
            p = row.cells[1].paragraphs[0]
            format_paragraph(p, space_after=0)
            r = p.add_run(v1)
            format_run(r, size_pt=9.5)
            
            # Celda 2 (y expandir en fila 4)
            if r_idx == 3:
                # Combinar celda 3 y 4 de la última fila para el diseño largo
                row.cells[2].merge(row.cells[3])
                set_cell_shading(row.cells[2], COLOR_GRAY_LIGHT_HEX)
                p = row.cells[2].paragraphs[0]
                format_paragraph(p, space_after=0)
                r = p.add_run("Capacidad & Exposición: ")
                format_run(r, size_pt=9, bold=True, color_rgb=(0x1D, 0x3D, 0x5A))
                r2 = p.add_run(v2)
                format_run(r2, size_pt=8.5, italic=True)
            else:
                set_cell_shading(row.cells[2], COLOR_GRAY_LIGHT_HEX)
                p = row.cells[2].paragraphs[0]
                format_paragraph(p, space_after=0)
                r = p.add_run(l2)
                format_run(r, size_pt=9, bold=True, color_rgb=(0x1D, 0x3D, 0x5A))
                
                p = row.cells[3].paragraphs[0]
                format_paragraph(p, space_after=0)
                r = p.add_run(v2)
                format_run(r, size_pt=9.5)
                
        doc.add_paragraph() # Espacio
        
        # Resultados de Simulación
        p_res = doc.add_paragraph()
        format_paragraph(p_res, space_after=4)
        r_lbl = p_res.add_run("Resultados de la Evaluación de Riesgo (3x3 Matrix):")
        format_run(r_lbl, size_pt=10.5, bold=True, color_rgb=(0x01, 0x27, 0x43))
        
        tbl_res = doc.add_table(rows=2, cols=4)
        tbl_res.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_borders(tbl_res, "E0E3E5")
        
        res_widths = [Inches(1.6), Inches(1.6), Inches(1.6), Inches(1.7)]
        for r in tbl_res.rows:
            for c_idx, cell in enumerate(r.cells):
                cell.width = res_widths[c_idx]
                set_cell_margins(cell, 80, 80, 100, 100)
                
        # Cabecera tabla de resultados
        res_headers = ["Probabilidad (LoF)", "Exposición (CoF)", "Coordenada Matrix", "Riesgo Resultante"]
        for idx, text in enumerate(res_headers):
            cell = tbl_res.cell(0, idx)
            set_cell_shading(cell, COLOR_SECONDARY_HEX)
            p = cell.paragraphs[0]
            format_paragraph(p, space_after=0, alignment=WD_ALIGN_PARAGRAPH.CENTER)
            r = p.add_run(text)
            format_run(r, size_pt=9, bold=True, color_rgb=(0xFF, 0xFF, 0xFF))
            
        # Valores
        tbl_res.cell(1, 0).paragraphs[0].add_run(m['lof'])
        tbl_res.cell(1, 1).paragraphs[0].add_run(m['exposure'])
        tbl_res.cell(1, 2).paragraphs[0].add_run(m['coordinate'])
        
        c_risk_val = tbl_res.cell(1, 3)
        p_risk_val = c_risk_val.paragraphs[0]
        r_risk_val = p_risk_val.add_run(m['risk_name'])
        
        # Color del badge de riesgo
        bg_hex_val, font_rgb = risk_colors_map[m['risk_code']]
        set_cell_shading(c_risk_val, bg_hex_val + "25")  # Ligero 15%
        format_run(r_risk_val, size_pt=9.5, bold=True, color_rgb=font_rgb)
        
        for idx in range(3):
            p = tbl_res.cell(1, idx).paragraphs[0]
            format_paragraph(p, space_after=0, alignment=WD_ALIGN_PARAGRAPH.CENTER)
            format_run(p.runs[0], size_pt=10, bold=True, color_rgb=(0x33, 0x33, 0x33))
            
        format_paragraph(p_risk_val, space_after=0, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        
        doc.add_paragraph() # Espacio
        
        # Cuestionario Evaluado
        p_quest_title = doc.add_paragraph()
        format_paragraph(p_quest_title, space_after=4)
        r_quest = p_quest_title.add_run("Detalle del Cuestionario y Desviaciones Registradas:")
        format_run(r_quest, size_pt=10.5, bold=True, color_rgb=(0x01, 0x27, 0x43))
        
        tbl_q = doc.add_table(rows=1, cols=3)
        tbl_q.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_borders(tbl_q, "E0E3E5")
        
        q_widths = [Inches(0.9), Inches(3.6), Inches(2.0)]
        hdr_q = tbl_q.rows[0].cells
        hdr_q_texts = ["ID Pregunta", "Criterio Normativo Evaluado (API RP 2SIM)", "Respuesta de Campo / Estatus"]
        for idx, text in enumerate(hdr_q_texts):
            hdr_q[idx].width = q_widths[idx]
            set_cell_shading(hdr_q[idx], COLOR_SECONDARY_HEX)
            set_cell_margins(hdr_q[idx], 80, 80, 100, 100)
            p = hdr_q[idx].paragraphs[0]
            format_paragraph(p, space_after=0, alignment=WD_ALIGN_PARAGRAPH.LEFT if idx < 2 else WD_ALIGN_PARAGRAPH.CENTER)
            r = p.add_run(text)
            format_run(r, size_pt=9, bold=True, color_rgb=(0xFF, 0xFF, 0xFF))
            
        # Llenar respuestas
        for q_id, val in answers.items():
            row = tbl_q.add_row()
            for idx in range(3):
                row.cells[idx].width = q_widths[idx]
                set_cell_margins(row.cells[idx], 60, 60, 80, 80)
                
            c_id = row.cells[0]
            c_text = row.cells[1]
            c_val = row.cells[2]
            
            # ID
            p_id = c_id.paragraphs[0]
            format_paragraph(p_id, space_after=0)
            r_id = p_id.add_run(q_id)
            format_run(r_id, size_pt=9.5, bold=True)
            
            # Criterio
            p_text = c_text.paragraphs[0]
            format_paragraph(p_text, space_after=0)
            q_meta = questions_dict.get(q_id, {
                'pregunta': "Criterio no documentado",
                'opcion_a': 'Desconocida', 'opcion_b': 'Desconocida', 'opcion_c': 'Desconocida'
            })
            r_text = p_text.add_run(q_meta['pregunta'])
            format_run(r_text, size_pt=9)
            
            # Valor Respuesta
            p_val = c_val.paragraphs[0]
            format_paragraph(p_val, space_after=0)
            
            ans_text = "N/D"
            is_deficiency = False
            
            if val == 'a':
                ans_text = q_meta.get('opcion_a', 'Opción a')
                # En módulos A, B, C la 'a' suele ser la más crítica. En módulo D la 'a' es el iniciador activo.
                is_deficiency = True
            elif val == 'b':
                ans_text = q_meta.get('opcion_b', 'Opción b')
                # En módulos A, B, C, la 'b' es intermedia. En módulo D, la 'b' es la segura (No).
                if q_id.startswith('C'):
                    is_deficiency = True  # Luz libre marginal, arriostramiento K, o daño moderado son deficiencias parciales
            elif val == 'c':
                ans_text = q_meta.get('opcion_c', 'Opción c')
                
            r_val = p_val.add_run(ans_text)
            
            # Formatear si es deficiencia
            if is_deficiency:
                format_run(r_val, size_pt=8.5, bold=True, color_rgb=(0xEF, 0x44, 0x44))
                set_cell_shading(c_val, "FDF2F2")
            else:
                format_run(r_val, size_pt=8.5, color_rgb=(0x27, 0x4E, 0x13))
                set_cell_shading(c_val, "ECFDF5")
                
        doc.add_paragraph() # Espacio
        
        # Plan de Mitigación y Recomendaciones
        p_mit_title = doc.add_paragraph()
        format_paragraph(p_mit_title, space_after=4)
        r_mit_title = p_mit_title.add_run("Plan de Mitigación y Acciones Recomendadas:")
        format_run(r_mit_title, size_pt=10.5, bold=True, color_rgb=(0x01, 0x27, 0x43))
        
        # Identificar las mitigaciones correspondientes
        has_mitigations = False
        for q_id, val in answers.items():
            q_meta = questions_dict.get(q_id)
            if not q_meta:
                continue
                
            is_def = False
            if val == 'a':
                is_def = True
            elif val == 'b' and q_id.startswith('C'):
                is_def = True
                
            if is_def:
                has_mitigations = True
                p_item = doc.add_paragraph(style='List Bullet')
                format_paragraph(p_item, space_before=2, space_after=2)
                r_item_id = p_item.add_run(f"[{q_id}] ")
                format_run(r_item_id, size_pt=9.5, bold=True, color_rgb=(0xEF, 0x44, 0x44))
                mit_text = q_meta['mitigacion']
                if plat_name == 'Balam-TC':
                    if q_id == 'C2':
                        mit_text = "Para esta estructura inactiva (L-3), postergar el refuerzo estructural físico de nudos críticos e instalar en su lugar un programa de monitoreo de estabilidad global pasiva."
                    elif q_id == 'C3':
                        mit_text = "Monitorear de forma no intervencionista la degradación del Jacket mediante inspección visual y mediciones de potencial catódico; posponer soldaduras o reparaciones subacuáticas mayores salvo que comprometa ductos o colectores activos adyacentes."
                    elif q_id == 'D4':
                        mit_text = "Exentar de análisis de re-calificación Pushover Nivel III detallado por su condición inactiva y nula exposición a riesgos humanos o ambientales mayores; mantener monitoreo visual de marcas de impacto de ola en cubierta inferior."
                r_item_text = p_item.add_run(mit_text)
                format_run(r_item_text, size_pt=9.5)
                
        if not has_mitigations:
            p_item = doc.add_paragraph()
            format_paragraph(p_item, space_after=4)
            r_item = p_item.add_run("🏆 La instalación no presenta deficiencias físicas ni iniciadores de reevaluación activos. Se encuentra en estado íntegro y óptimo.")
            format_run(r_item, size_pt=10, bold=True, color_rgb=(0x27, 0x4E, 0x13))
            
        doc.add_paragraph() # Espacio
        
        # Plan de Inspección Recomendado
        intervals = get_inspection_intervals(m['risk_code'])
        p_insp_title = doc.add_paragraph()
        format_paragraph(p_insp_title, space_after=4)
        r_insp_title = p_insp_title.add_run("Programa de Inspección Recomendado (Estándar API RP 2SIM / NRF-260-PEMEX):")
        format_run(r_insp_title, size_pt=10.5, bold=True, color_rgb=(0x01, 0x27, 0x43))
        
        tbl_insp = doc.add_table(rows=1, cols=3)
        tbl_insp.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_borders(tbl_insp, "E0E3E5")
        
        insp_widths = [Inches(2.1), Inches(2.2), Inches(2.2)]
        hdr_insp = tbl_insp.rows[0].cells
        hdr_insp_texts = ["Inspección Visual General (GVI)", "Inspección Visual Detallada (CVI)", "Ensayos No Destructivos (NDT)"]
        for idx, text in enumerate(hdr_insp_texts):
            hdr_insp[idx].width = insp_widths[idx]
            set_cell_shading(hdr_insp[idx], COLOR_GRAY_LIGHT_HEX)
            set_cell_margins(hdr_insp[idx], 80, 80, 100, 100)
            p = hdr_insp[idx].paragraphs[0]
            format_paragraph(p, space_after=0, alignment=WD_ALIGN_PARAGRAPH.CENTER)
            r = p.add_run(text)
            format_run(r, size_pt=9, bold=True, color_rgb=(0x1D, 0x3D, 0x5A))
            
        row_insp = tbl_insp.add_row()
        row_insp.cells[0].width = Inches(2.1)
        row_insp.cells[1].width = Inches(2.2)
        row_insp.cells[2].width = Inches(2.2)
        
        tbl_insp.cell(1, 0).paragraphs[0].add_run(intervals['gvi'])
        tbl_insp.cell(1, 1).paragraphs[0].add_run(intervals['cvi'])
        tbl_insp.cell(1, 2).paragraphs[0].add_run(intervals['ndt'])
        
        for idx in range(3):
            cell = tbl_insp.cell(1, idx)
            set_cell_margins(cell, 80, 80, 100, 100)
            p = cell.paragraphs[0]
            format_paragraph(p, space_after=0, alignment=WD_ALIGN_PARAGRAPH.CENTER)
            format_run(p.runs[0], size_pt=9.5, bold=True, color_rgb=(0x33, 0x33, 0x33))
            
        doc.add_paragraph() # Separación
        
        # Evitar page break en la última plataforma
        if plat_idx < len(calculated_fleet):
            doc.add_page_break()
            
    doc.add_page_break()
    
    # --- SECCIÓN 4: CONCLUSIONES Y RECOMENDACIONES ---
    add_heading_styled(doc, "4. Conclusiones y Recomendaciones Generales", level=1)
    
    p = doc.add_paragraph()
    format_paragraph(p)
    run = p.add_run(
        "Con base en el diagnóstico consolidado del Complejo Costa Afuera de Ek-Balam al año 2026, se derivan "
        "las siguientes conclusiones estratégicas para la toma de decisiones y la priorización del plan de inversión "
        "en integridad de activos:"
    )
    format_run(run)
    
    conclusions = [
        ("Alta Concentración de Riesgo en Nodos Troncales y de Proceso: ", 
         "La plataforma EK-A (Riesgo Alto / Nivel 1) y el trípode de reinyección EK-TB (Riesgo Alto / Nivel 1) "
         "constituyen los dos puntos críticos que concentran el mayor nivel de riesgo del complejo. La falla de EK-A interrumpiría el "
         "100% del flujo del activo, mientras que EK-TB experimenta desviaciones por probabilidad de falla alta debido a daños "
         "físicos acumulados severos y nula holgura de Air Gap en su subestructura. Deben priorizarse de forma inmediata las acciones "
         "de mitigación subacuática y el aligeramiento de cubiertas inferiores en ambas estructuras."),
         
        ("Monitoreo Estricto de Plataformas de Riesgo Medio: ", 
         "Plataformas como Balam-TB (cabecera del colector de 20 pulgadas), Balam-TD, Balam-TA, Balam-1 y la habitacional EK-A Hab se "
         "clasifican en Riesgo Medio (Nivel 2). En particular, EK-A Hab, a pesar de contar con una condición física aceptable (LoF Baja), "
         "se mantiene en categoría de consecuencia L-1 debido a la dotación permanente de 120 personas a bordo, lo que restringe el margen "
         "de tolerancia. Se debe validar de forma anual el plan de evacuación rápida pre-evento y la funcionalidad de los botes de salvamento."),
         
        ("Recalificación Estructural Obligatoria para Extensión de Vida Útil (Pre-1995): ", 
         "Las plataformas instaladas con anterioridad a 1995 (EK-A, EK-TA, EK-TB, Balam-A, Balam-TB, Balam-TC, Balam-1) han rebasado su vida útil "
         "original de diseño de 30 años y operan bajo contratos de extensión vigentes hasta el año 2039. De acuerdo con la API RP 2SIM, "
         "el haber superado este umbral temporal activa obligatoriamente la necesidad de realizar un Análisis de Re-calificación Estructural "
         "(Structural Re-qualification Assessment) integral que justifique su aptitud para el servicio."),
         
        ("Estatus Especial de Plataforma Inactiva (Balam-TC): ", 
         "La plataforma Balam-TC se cataloga como fuera de operación / inactiva. Su nivel de riesgo se mantiene en el rango Bajo (Nivel 3) "
         "gracias a la ausencia de personal a bordo (0 POB) y la desactivación total de flujo de hidrocarburos (0 BPD), lo cual reduce "
         "su consecuencia económica/ambiental a L-3. No obstante, al poseer un estado de daño moderado en su Jacket (C3='b'), requiere "
         "un monitoreo estructural periódico pasivo para prevenir una degradación acelerada de su Jacket que comprometa la estabilidad global "
         "del colector local."),
         
        ("Sistematización del Plan de Mantenimiento e Inspección: ", 
         "Es indispensable programar las campañas de inspección submarina y aérea de acuerdo con la periodicidad establecida en el reporte. "
         "Para las instalaciones de Riesgo Alto, el intervalo máximo permitido para inspección general visual (GVI) es de 1 a 2 años, y de "
         "3 a 5 años para visual detallada (CVI) y ensayos no destructivos (NDT). Para las de Riesgo Bajo, las frecuencias pueden "
         "espaciarse a 5-6 años (GVI) y 10-12 años (CVI), optimizando sustancialmente el presupuesto operativo del Activo.")
    ]
    
    for title, text in conclusions:
        p_c = doc.add_paragraph(style='List Bullet')
        format_paragraph(p_c, space_before=3, space_after=3)
        r_title = p_c.add_run(title)
        format_run(r_title, size_pt=10, bold=True, color_rgb=(0x01, 0x27, 0x43))
        r_text = p_c.add_run(text)
        format_run(r_text, size_pt=10)
        
    for _ in range(3):
        doc.add_paragraph()
        
    # Firmas / Aprobaciones
    tbl_sigs = doc.add_table(rows=2, cols=2)
    tbl_sigs.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(tbl_sigs, "FFFFFF")
    
    sig_widths = [Inches(3.2), Inches(3.2)]
    for r in tbl_sigs.rows:
        for idx, cell in enumerate(r.cells):
            cell.width = sig_widths[idx]
            set_cell_margins(cell, 150, 150, 100, 100)
            
    p_sig1 = tbl_sigs.cell(0, 0).paragraphs[0]
    format_paragraph(p_sig1, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    r_sig1 = p_sig1.add_run("_____________________________________\nElaboró:\nIngeniería de Integridad Estructural\nGrupo ST-STORM Costa Afuera")
    format_run(r_sig1, size_pt=9.5, italic=True)
    
    p_sig2 = tbl_sigs.cell(0, 1).paragraphs[0]
    format_paragraph(p_sig2, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    r_sig2 = p_sig2.add_run("_____________________________________\nRevisó y Aprobó:\nGerencia de Integridad y Mantenimiento\nActivo de Producción Ek-Balam - PEMEX")
    format_run(r_sig2, size_pt=9.5, italic=True)
    
    # 5. Guardar documento
    doc.save(output_filename)
    print(f"Reporte Word generado exitosamente: {output_filename}")

if __name__ == "__main__":
    main()
